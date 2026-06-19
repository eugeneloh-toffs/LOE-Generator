"""
Generate a Letter of Employment .docx template with highlighted {{ placeholders }}.

Every placeholder is emitted as its OWN run with a yellow highlight, so:
  - a human can see exactly what gets replaced, and
  - docxtpl still detects/renders it (placeholder is never split across runs).

Run:  python make_template.py
Out:  samples/LOE_template.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor

PLACEHOLDER_RE = re.compile(r"(\{\{.*?\}\})")


def add_runs(paragraph, text, *, bold=False, italic=False, size=None):
    """Split text on {{ placeholders }}; highlight placeholder runs in yellow."""
    for part in PLACEHOLDER_RE.split(text):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if part.startswith("{{"):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return paragraph


def para(doc, text="", *, bold=False, italic=False, size=None, align=None,
         space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    add_runs(p, text, bold=bold, italic=italic, size=size)
    return p


def heading(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{number}. {text}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)  # navy
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_runs(p, text)
    return p


doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ----------------------------------------------------------------------------
# Letterhead
# ----------------------------------------------------------------------------
lh = doc.add_paragraph()
lh.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = lh.add_run("{{ company_name }}")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
r.font.highlight_color = WD_COLOR_INDEX.YELLOW

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
add_runs(sub, "{{ company_address }}", size=9)
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(12)
add_runs(contact, "UEN: {{ company_uen }}  |  Tel: {{ company_phone }}  |  {{ company_email }}", size=9)

# Date + reference
para(doc, "Date: {{ offer_date }}")
para(doc, "Our Ref: {{ reference_number }}", space_after=12)

# Recipient block
para(doc, "{{ name }}", bold=True, space_after=2)
para(doc, "{{ candidate_address }}", space_after=2)
para(doc, "NRIC / FIN: {{ nric }}", space_after=12)

# Salutation + subject
para(doc, "Dear {{ first_name }},", space_after=8)
subj = doc.add_paragraph()
subj.paragraph_format.space_after = Pt(10)
rs = subj.add_run("RE: LETTER OF EMPLOYMENT — ")
rs.bold = True
rs2 = subj.add_run("{{ job_title }}")
rs2.bold = True
rs2.font.highlight_color = WD_COLOR_INDEX.YELLOW

para(doc,
     "We are pleased to offer you employment with {{ company_name }} (the “Company”) "
     "on the terms and conditions set out in this letter. This letter, together with the "
     "Company’s Employee Handbook and any policies referenced herein, forms the basis of "
     "your employment.", space_after=10)

# ----------------------------------------------------------------------------
# 1. Appointment & Reporting
# ----------------------------------------------------------------------------
heading(doc, 1, "Appointment and Reporting")
para(doc,
     "You are appointed as {{ job_title }} in the {{ department }} department. You will "
     "report directly to {{ reporting_manager }} ({{ reporting_manager_title }}), or to "
     "such other person as the Company may designate from time to time.")
para(doc,
     "Your principal place of work will be {{ work_location }}. The Company may require you "
     "to work at, or be transferred to, other locations as reasonably necessary for the "
     "performance of your duties.")
para(doc,
     "A summary of your key responsibilities is set out below. This list is not exhaustive, "
     "and you may be required to perform other duties commensurate with your role:")
bullet(doc, "Deliver the core functions of the {{ job_title }} role to the standards set by the Company;")
bullet(doc, "Comply with all lawful and reasonable instructions given by {{ reporting_manager }};")
bullet(doc, "Act in the best interests of the Company at all times; and")
bullet(doc, "Observe all Company policies, procedures, and codes of conduct in force.")

# ----------------------------------------------------------------------------
# 2. Commencement & Probation
# ----------------------------------------------------------------------------
heading(doc, 2, "Commencement Date and Probation")
para(doc,
     "Your employment will commence on {{ start_date }}, subject to the Company’s receipt "
     "of all documents and clearances referred to in Clause 12. The first {{ probation_period }} "
     "of your employment shall be a probationary period.")
para(doc,
     "During the probationary period, either party may terminate the employment by giving "
     "{{ probation_notice }} written notice, or payment in lieu thereof. The Company may, at "
     "its sole discretion, extend the probationary period by up to {{ probation_extension }} "
     "where it considers further assessment appropriate. Upon successful completion of "
     "probation, your employment will be confirmed in writing.")

# ----------------------------------------------------------------------------
# 3. Remuneration
# ----------------------------------------------------------------------------
heading(doc, 3, "Remuneration")
para(doc,
     "Your gross monthly salary will be SGD {{ salary }}, payable monthly in arrears on or "
     "before {{ payment_date }} of each calendar month by direct credit to your nominated "
     "bank account. Your salary will be reviewed annually, with any adjustment being at the "
     "sole discretion of the Company and not guaranteed.")
para(doc,
     "You may be eligible for a discretionary performance bonus of up to {{ bonus_target }} "
     "of your annual base salary, subject to your individual performance and the Company’s "
     "overall performance. Any bonus is entirely discretionary and is not payable if you are "
     "under notice of termination (whether given or received) at the payment date.")
para(doc,
     "The Company will make contributions to the Central Provident Fund (CPF) in accordance "
     "with the Central Provident Fund Act, and will deduct your employee CPF contributions "
     "from your salary as required by law.")

# ----------------------------------------------------------------------------
# 4. Working Hours
# ----------------------------------------------------------------------------
heading(doc, 4, "Working Hours")
para(doc,
     "Your normal working hours are {{ working_hours }}, {{ working_days }}, with a "
     "{{ lunch_break }} lunch break. Given the nature of your role, you may be required to "
     "work additional hours as reasonably necessary to fulfil your duties, for which no "
     "additional overtime payment will be made unless required by the Employment Act.")

# ----------------------------------------------------------------------------
# 5. Leave & Benefits
# ----------------------------------------------------------------------------
heading(doc, 5, "Leave and Benefits")
para(doc, "You will be entitled to the following leave and benefits, subject to the Company’s policies:")
bullet(doc, "Annual leave: {{ annual_leave }} working days per calendar year, accruing monthly;")
bullet(doc, "Paid sick leave: {{ medical_leave }} days of outpatient sick leave and {{ hospitalisation_leave }} days of hospitalisation leave per year, supported by a valid medical certificate;")
bullet(doc, "Medical benefits: {{ medical_benefits }};")
bullet(doc, "Insurance coverage: {{ insurance_coverage }};")
bullet(doc, "Other benefits: {{ other_benefits }}.")
para(doc,
     "Unused annual leave may be carried forward only to the extent permitted by the "
     "Company’s leave policy. All benefits are subject to the terms of the relevant "
     "policy or insurance scheme and may be varied by the Company from time to time.")

# ----------------------------------------------------------------------------
# 6. Confidentiality
# ----------------------------------------------------------------------------
heading(doc, 6, "Confidentiality")
para(doc,
     "During and after your employment, you shall keep confidential and shall not, without "
     "the prior written consent of the Company, disclose to any third party or use for your "
     "own benefit any confidential information belonging to the Company, its clients, or its "
     "business partners. Confidential information includes, but is not limited to, business "
     "plans, client lists, pricing, technical data, source code, and trade secrets.")
para(doc,
     "All documents, records, and materials created or received by you in the course of your "
     "employment remain the property of the Company and must be returned upon termination of "
     "your employment or earlier upon request.")

# ----------------------------------------------------------------------------
# 7. Intellectual Property
# ----------------------------------------------------------------------------
heading(doc, 7, "Intellectual Property")
para(doc,
     "All intellectual property, including inventions, designs, works, and improvements, "
     "created by you in the course of your employment shall belong absolutely to the Company. "
     "You agree to execute all documents and do all things reasonably necessary to vest such "
     "rights in the Company, and you waive any moral rights you may have in such works to the "
     "extent permitted by law.")

# ----------------------------------------------------------------------------
# 8. Restrictive Covenants
# ----------------------------------------------------------------------------
heading(doc, 8, "Restrictive Covenants")
para(doc,
     "For a period of {{ non_compete_period }} following the termination of your employment, "
     "you shall not, within {{ non_compete_geography }}, directly or indirectly engage in any "
     "business that competes with the Company, nor solicit any client or employee of the "
     "Company with whom you had material dealings during the {{ lookback_period }} preceding "
     "termination. The parties agree that these restrictions are reasonable and necessary to "
     "protect the legitimate business interests of the Company.")

# ----------------------------------------------------------------------------
# 9. Code of Conduct
# ----------------------------------------------------------------------------
heading(doc, 9, "Code of Conduct and Outside Activities")
para(doc,
     "You shall devote your full working time and attention to your duties and shall not, "
     "without the prior written consent of the Company, engage in any other business or "
     "employment during the term of your employment. You shall comply with all applicable "
     "laws and the Company’s policies on anti-bribery, conflicts of interest, and "
     "acceptable use of Company systems.")

# ----------------------------------------------------------------------------
# 10. Termination
# ----------------------------------------------------------------------------
heading(doc, 10, "Termination of Employment")
para(doc,
     "After confirmation, either party may terminate this employment by giving "
     "{{ notice_period }} written notice, or salary in lieu of notice. The Company reserves "
     "the right to terminate your employment summarily without notice or payment in lieu in "
     "cases of misconduct, breach of this letter, or other circumstances justifying summary "
     "dismissal under the Employment Act.")
para(doc,
     "Upon termination, you shall return all Company property and shall be paid all salary "
     "and accrued but unused annual leave up to your last day of service, less any sums owed "
     "by you to the Company.")

# ----------------------------------------------------------------------------
# 11. Data Protection
# ----------------------------------------------------------------------------
heading(doc, 11, "Personal Data")
para(doc,
     "By accepting this offer, you consent to the Company collecting, using, and disclosing "
     "your personal data for purposes reasonably required for the management of your "
     "employment, in accordance with the Personal Data Protection Act 2012 and the Company’s "
     "data protection policy.")

# ----------------------------------------------------------------------------
# 12. Conditions of Offer
# ----------------------------------------------------------------------------
heading(doc, 12, "Conditions of Offer")
para(doc, "This offer of employment is conditional upon:")
bullet(doc, "Verification of your identity and right to work in Singapore;")
bullet(doc, "Satisfactory reference and background checks;")
bullet(doc, "Your possession of a valid {{ work_pass_type }}, where applicable; and")
bullet(doc, "Your signing and returning this letter by {{ acceptance_deadline }}.")

# ----------------------------------------------------------------------------
# 13. Medical Examination
# ----------------------------------------------------------------------------
heading(doc, 13, "Medical Examination and Fitness to Work")
para(doc,
     "The Company may require you to undergo a pre-employment medical examination by a "
     "medical practitioner nominated by the Company, the satisfactory outcome of which is a "
     "condition of this offer. During your employment, the Company may also require you to "
     "undergo such further medical examinations as may be reasonably necessary, the cost of "
     "which shall be borne by the Company. You agree to authorise the disclosure of the "
     "results of any such examination to the Company for the purpose of assessing your "
     "fitness to perform the duties of {{ job_title }}.")

# ----------------------------------------------------------------------------
# 14. Variation
# ----------------------------------------------------------------------------
heading(doc, 14, "Variation of Terms")
para(doc,
     "The Company reserves the right to make reasonable changes to any of the terms and "
     "conditions of your employment, including your duties, reporting line, place of work, "
     "and benefits, to meet the changing needs of the business. Any material variation to "
     "these terms will be communicated to you in writing and will take effect on the date "
     "specified in that notice. Your continued performance of your duties after such notice "
     "shall constitute your acceptance of the varied terms.")

# ----------------------------------------------------------------------------
# 15. Governing Law
# ----------------------------------------------------------------------------
heading(doc, 15, "Entire Agreement and Governing Law")
para(doc,
     "This letter constitutes the entire agreement between the parties and supersedes all "
     "prior discussions and representations. It shall be governed by and construed in "
     "accordance with the laws of Singapore, and the parties submit to the exclusive "
     "jurisdiction of the Singapore courts.")

# ----------------------------------------------------------------------------
# Acceptance & signatures
# ----------------------------------------------------------------------------
para(doc,
     "We look forward to welcoming you to {{ company_name }} and are confident that you will "
     "make a valuable contribution to our team. Please confirm your acceptance of this offer "
     "by signing and returning a copy of this letter.", space_before=10, space_after=16)

para(doc, "Yours sincerely,", space_after=24)
para(doc, "{{ hr_name }}", bold=True, space_after=2)
para(doc, "{{ hr_title }}", space_after=2)
para(doc, "{{ company_name }}", space_after=20)

# Acceptance block
acc = doc.add_paragraph()
acc.paragraph_format.space_before = Pt(10)
acc.paragraph_format.space_after = Pt(10)
ra = acc.add_run("ACKNOWLEDGEMENT AND ACCEPTANCE")
ra.bold = True
ra.font.size = Pt(12)
ra.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

para(doc,
     "I, {{ name }} (NRIC/FIN: {{ nric }}), have read, understood, and accept the terms and "
     "conditions of employment set out in this letter.", space_after=30)

para(doc, "Signature: ____________________________", space_after=18)
para(doc, "Name: {{ name }}", space_after=18)
para(doc, "Date: ____________________________", space_after=6)

# ----------------------------------------------------------------------------
out = Path(__file__).parent / "samples" / "LOE_template.docx"
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)
print(f"Saved: {out}")

# Report detected placeholders
all_text = "\n".join(p.text for p in doc.paragraphs)
found = sorted(set(re.findall(r"\{\{\s*(\w+)\s*\}\}", all_text)))
print(f"\n{len(found)} unique placeholders (use these as Excel column headers):")
for name in found:
    print(f"  - {name}")
